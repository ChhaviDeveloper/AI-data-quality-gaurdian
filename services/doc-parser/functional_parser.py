# functional_parser.py
"""
Parse a functional Word document (.docx) and generate a rules_registry.yaml.

Copied from the root-level python_scripts/functional_parser.py so this
Cloud Run service is self-contained (no dependency on the rest of the repo
at build time). Keep this in sync with the root copy, or replace both with
a shared package later if the two start to drift.

Usage:
    - Local docx: generate_rules_registry_from_functional_doc("path/to/Functional_Document.docx", "specs/rules_registry.yaml")
    - GCS docx:  generate_rules_registry_from_functional_doc("gs://bucket/path/Functional_Document.docx", "gs://bucket/path/specs/rules_registry.yaml")
"""

import os
import re
import yaml
from io import BytesIO

try:
    from docx import Document
except Exception:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

# Optional GCS support
GCS_PREFIX = "gs://"
def _read_docx_text_from_gcs(gcs_path):
    from google.cloud import storage
    parsed = gcs_path[len(GCS_PREFIX):].split("/", 1)
    bucket_name, blob_path = parsed[0], parsed[1]
    client = storage.Client()
    bucket = client.bucket(bucket_name)
    blob = bucket.blob(blob_path)
    data = blob.download_as_bytes()
    return BytesIO(data)

def _load_docx(path):
    """
    Returns a python-docx Document object.
    Accepts local path or gs:// path.
    """
    if path.startswith(GCS_PREFIX):
        bio = _read_docx_text_from_gcs(path)
        return Document(bio)
    else:
        return Document(path)

# Mapping of rule IDs to expressions and types (seeded from your functional doc)
# Keep complex checks as type: python (use RuleContext helpers in validator)
RULE_TEMPLATES = {
    "R001": {"type":"expression", "expression":"application_id == '' or application_id.isnull()"},
    "R002": {"type":"python", "expression":"ctx.duplicated_mask('application_id')"},
    "R003": {"type":"expression", "expression":"app_status == 'Active' and (business_owner == '' or technology_owner == '')"},
    "R004": {"type":"expression", "expression":"criticality not in ['Critical','High','Medium','Low']"},
    "R005": {"type":"python", "expression":"ctx.invalid_yes_no_mask(['evidence_submitted','access_review_completed','privileged_access_approved','mfa_enabled','encryption_at_rest','backup_enabled','policy_exception'])"},
    "R006": {"type":"expression", "expression":"regulatory_scope in ['SOX','RBI','SEBI','Internal Audit'] and evidence_submitted != 'Yes'"},
    "R007": {"type":"expression", "expression":"criticality in ['Critical','High'] and access_review_completed != 'Yes'"},
    "R008": {"type":"expression", "expression":"privileged_access_count > 0 and privileged_access_approved != 'Yes'"},
    "R009": {"type":"expression", "expression":"criticality in ['Critical','High'] and mfa_enabled != 'Yes'"},
    "R010": {"type":"expression", "expression":"data_classification in ['Confidential','Restricted'] and encryption_at_rest != 'Yes'"},
    "R011": {"type":"expression", "expression":"app_status == 'Active' and criticality in ['Critical','High'] and backup_enabled != 'Yes'"},
    "R012": {"type":"python", "expression":"ctx.dr_test_invalid_mask()"},
    "R013": {"type":"expression", "expression":"(vulnerability_status == 'Critical Open' or open_high_vulnerabilities > 0) and policy_exception != 'Yes'"},
    "R014": {"type":"python", "expression":"ctx.exception_expiry_invalid_mask()"},
    "R015": {"type":"python", "expression":"ctx.future_dr_dates()"}
}

def _extract_rules_table(doc):
    """
    Heuristic: find the table that contains a header with 'Rule ID' or 'Rule Name' or 'Business Rule / Validation Logic'.
    Return list of rows as lists of cell texts.
    """
    candidate_tables = []
    for table in doc.tables:
        # collect header text
        header_text = " ".join([cell.text.strip() for cell in table.rows[0].cells]).lower()
        if any(k in header_text for k in ("rule id", "rule name", "business rule", "validation logic")):
            candidate_tables.append(table)
    if not candidate_tables:
        # fallback: return the largest table
        tables_sorted = sorted(doc.tables, key=lambda t: len(t.rows), reverse=True)
        if tables_sorted:
            return tables_sorted[0]
        return None
    # prefer the first candidate
    return candidate_tables[0]

def _normalize_cell_text(text):
    if text is None:
        return ""
    return re.sub(r"\s+", " ", text).strip()

def _table_to_rule_dicts(table):
    """
    Convert a docx table into a list of rule dicts.
    Expects columns: Rule ID | Rule Name | Business Rule / Validation Logic | Severity | Data Quality Dimension
    If columns differ, attempt to map by header keywords.
    """
    rows = list(table.rows)
    if not rows or len(rows) < 2:
        return []

    # build header mapping
    headers = [ _normalize_cell_text(cell.text).lower() for cell in rows[0].cells ]
    # map header index by keyword
    idx_map = {}
    for i,h in enumerate(headers):
        if "rule id" in h or re.search(r"\brule\b.*\bid\b", h):
            idx_map["rule_id"] = i
        elif "rule name" in h or "name" in h:
            idx_map["rule_name"] = i
        elif "business rule" in h or "validation logic" in h or "logic" in h:
            idx_map["logic"] = i
        elif "severity" in h:
            idx_map["severity"] = i
        elif "dimension" in h or "data quality" in h:
            idx_map["dimension"] = i
    # fallback heuristics
    if "rule_id" not in idx_map and len(headers) >= 1:
        idx_map["rule_id"] = 0
    if "rule_name" not in idx_map and len(headers) >= 2:
        idx_map["rule_name"] = 1
    if "logic" not in idx_map and len(headers) >= 3:
        idx_map["logic"] = 2
    if "severity" not in idx_map and len(headers) >= 4:
        idx_map["severity"] = 3
    if "dimension" not in idx_map and len(headers) >= 5:
        idx_map["dimension"] = 4

    rule_dicts = []
    for r in rows[1:]:
        cells = [ _normalize_cell_text(c.text) for c in r.cells ]
        # guard: skip empty rows
        if all(not c for c in cells):
            continue
        rule_id = cells[idx_map.get("rule_id","")] if idx_map.get("rule_id","") != "" else ""
        rule_name = cells[idx_map.get("rule_name","")] if idx_map.get("rule_name","") != "" else ""
        logic = cells[idx_map.get("logic","")] if idx_map.get("logic","") != "" else ""
        severity = cells[idx_map.get("severity","")] if idx_map.get("severity","") != "" else ""
        dimension = cells[idx_map.get("dimension","")] if idx_map.get("dimension","") != "" else ""
        rule_dicts.append({
            "rule_id": rule_id.strip(),
            "rule_name": rule_name.strip(),
            "logic": logic.strip(),
            "severity": severity.strip(),
            "dimension": dimension.strip()
        })
    return rule_dicts

def _map_to_registry_entry(rule_row):
    """
    Map a parsed rule row to a registry entry.
    If rule_id is known (R001..R015) use the template expression/type.
    Otherwise create a safe placeholder expression and mark type=expression.
    """
    rid = rule_row.get("rule_id","").upper()
    name = rule_row.get("rule_name") or f"Rule {rid or 'Unknown'}"
    desc = rule_row.get("logic") or rule_row.get("rule_name") or ""
    severity = rule_row.get("severity") or "High"
    dimension = rule_row.get("dimension") or "Unknown"

    if rid in RULE_TEMPLATES:
        tpl = RULE_TEMPLATES[rid]
        entry = {
            "rule_id": rid,
            "rule_name": name,
            "severity": severity,
            "dimension": dimension,
            "type": tpl.get("type","expression"),
            "expression": tpl.get("expression",""),
            "description": desc
        }
    else:
        # create a safe placeholder expression that will not error in pandas.query
        # we store the original logic in description for manual review
        placeholder_expr = "False"  # no rows will match; forces human review before enabling
        entry = {
            "rule_id": rid or "",
            "rule_name": name,
            "severity": severity,
            "dimension": dimension,
            "type": "expression",
            "expression": placeholder_expr,
            "description": desc or "Auto-generated placeholder. Please update expression."
        }
    return entry

def generate_rules_registry_from_functional_doc(functional_doc_path, output_registry_path):
    """
    Main entrypoint.
    - functional_doc_path: local path or gs:// path to .docx
    - output_registry_path: local path or gs:// path to write rules_registry.yaml
    """
    doc = _load_docx(functional_doc_path)
    table = _extract_rules_table(doc)
    if table is None:
        raise RuntimeError("Could not find rules table in the functional document. Ensure the document contains the 'Business Rules and Data Quality Logic' table.")

    parsed_rows = _table_to_rule_dicts(table)
    if not parsed_rows:
        raise RuntimeError("No rule rows extracted from the functional document table.")

    registry_rules = []
    for row in parsed_rows:
        entry = _map_to_registry_entry(row)
        registry_rules.append(entry)

    registry = {"rules": registry_rules}

    # Save registry to local or GCS
    if output_registry_path.startswith(GCS_PREFIX):
        from google.cloud import storage
        parsed = output_registry_path[len(GCS_PREFIX):].split("/",1)
        bucket_name, blob_path = parsed[0], parsed[1]
        client = storage.Client()
        bucket = client.bucket(bucket_name)
        blob = bucket.blob(blob_path)
        blob.upload_from_string(yaml.safe_dump(registry, sort_keys=False), content_type="text/yaml")
    else:
        os.makedirs(os.path.dirname(output_registry_path) or ".", exist_ok=True)
        with open(output_registry_path, "w", encoding="utf-8") as f:
            yaml.safe_dump(registry, f, sort_keys=False)

    return registry

# If run as script, quick smoke test (local)
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate rules_registry.yaml from functional .docx")
    parser.add_argument("--input", "-i", required=True, help="Path to functional docx (local or gs://)")
    parser.add_argument("--output", "-o", default="specs/rules_registry.yaml", help="Output registry path (local or gs://)")
    args = parser.parse_args()
    reg = generate_rules_registry_from_functional_doc(args.input, args.output)
    print(f"Wrote {len(reg['rules'])} rules to {args.output}")
